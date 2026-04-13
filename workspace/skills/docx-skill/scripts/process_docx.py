#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx-skill: Extract text from .docx file
Usage: python process_docx.py <input.docx>
"""
import sys
from docx import Document

if len(sys.argv) != 2:
    print("Usage: python process_docx.py <input.docx>")
    sys.exit(1)

file_path = sys.argv[1]
try:
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)
    print("\n".join(full_text))
except Exception as e:
    print(f"Error reading {file_path}: {e}")
    sys.exit(1)
