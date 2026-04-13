#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extract text from .docx file and save to .txt
"""
import sys
import os
from docx import Document

if len(sys.argv) != 2:
    print("Usage: python extract_text.py <input.docx>")
    sys.exit(1)

file_path = sys.argv[1]
try:
    # Try to read the document
    doc = Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
    
    # Create output filename
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_file = os.path.join(os.path.dirname(file_path), f"{base_name}_text.txt")
    
    # Write to txt file with UTF-8 encoding
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("\n".join(full_text))
    
    print(f"Text extracted successfully to: {output_file}")
    
except Exception as e:
    print(f"Error reading {file_path}: {e}")
    sys.exit(1)
